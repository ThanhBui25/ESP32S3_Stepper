import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(15)
    r.bold = True
    r.font.color.rgb = RGBColor(30, 58, 138) # Dark Navy

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(29, 78, 216) # Royal Blue

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 41, 55)

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.font.name = "Segoe UI"
        r_b.font.size = Pt(10)
        r_b.bold = True
        r_b.font.color.rgb = RGBColor(31, 41, 55)
    r_t = p.add_run(text)
    r_t.font.name = "Segoe UI"
    r_t.font.size = Pt(10)
    r_t.font.color.rgb = RGBColor(55, 65, 81)

def add_paragraph(doc, text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.font.name = "Segoe UI"
        r_b.font.size = Pt(10.5)
        r_b.bold = True
        r_b.font.color.rgb = RGBColor(31, 41, 55)
    r_t = p.add_run(text)
    r_t.font.name = "Segoe UI"
    r_t.font.size = Pt(10.5)
    r_t.font.color.rgb = RGBColor(55, 65, 81)

def add_alert_box(doc, title, text, alert_type="warning"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    
    if alert_type == "danger":
        bg_color = "FEF2F2"
        border_color = "DC2626"
        t_color = RGBColor(220, 38, 38)
    elif alert_type == "warning":
        bg_color = "FFFBEB"
        border_color = "D97706"
        t_color = RGBColor(217, 119, 6)
    else: # note/info
        bg_color = "EFF6FF"
        border_color = "2563EB"
        t_color = RGBColor(37, 99, 235)
        
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=100, bottom=100, left=160, right=140)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'  <w:top w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    r_title = p.add_run(f"{title}\n")
    r_title.font.name = "Segoe UI"
    r_title.font.size = Pt(10)
    r_title.bold = True
    r_title.font.color.rgb = t_color
    
    r_text = p.add_run(text)
    r_text.font.name = "Segoe UI"
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(31, 41, 55)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_table_data(doc, headers, data, col_widths=None):
    num_cols = len(headers)
    table = doc.add_table(rows=len(data) + 1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    
    if not col_widths:
        col_width = Inches(6.8 / num_cols)
        col_widths = [col_width] * num_cols
    else:
        col_widths = [Inches(w) for w in col_widths]
        
    # Header row
    hdr_row = table.rows[0]
    for idx, h_text in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h_text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(9.5)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    # Data rows
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        is_even = (r_idx % 2 == 0)
        bg = "F8FAFC" if is_even else "FFFFFF"
        
        for c_idx, cell_value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            
            # Format text in cell
            r = p.add_run(cell_value)
            r.font.name = "Segoe UI"
            r.font.size = Pt(9)
            if c_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            else:
                r.font.color.rgb = RGBColor(51, 65, 85)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def generate_full_hardware_docx(docx_path):
    doc = docx.Document()
    
    # Page setup
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
        footer = s.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Tài liệu Hướng dẫn Phần cứng & Sơ đồ chân Thiết bị - Dự án Điều khiển Tự động")
        f_run.font.name = "Segoe UI"
        f_run.font.size = Pt(8.5)
        f_run.font.color.rgb = RGBColor(148, 163, 184)

    # Document Header
    p_top = doc.add_paragraph()
    p_top.paragraph_format.space_before = Pt(10)
    p_top.paragraph_format.space_after = Pt(2)
    r_top = p_top.add_run("HỆ THỐNG ĐIỀU KHIỂN & TỰ ĐỘNG HÓA")
    r_top.font.name = "Segoe UI"
    r_top.font.size = Pt(11)
    r_top.bold = True
    r_top.font.color.rgb = RGBColor(37, 99, 235)
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after = Pt(8)
    r_title = p_title.add_run("SỔ TAY TRA CỨU PHẦN CỨNG CHI TIẾT\n(Chức năng - Sơ đồ chân - Lưu ý kỹ thuật - Cảnh báo)")
    r_title.font.name = "Segoe UI"
    r_title.font.size = Pt(18)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(30, 58, 138)
    
    add_alert_box(
        doc,
        "📌 MỤC ĐÍCH TÀI LIỆU:",
        "Tài liệu kỹ thuật tổng hợp toàn diện chức năng, thông số, sơ đồ chân chi tiết từng cổng và các lưu ý an toàn tối quan trọng của 11 thiết bị/linh kiện cốt lõi trong hệ thống điều khiển tự động hóa.",
        "info"
    )

    # =========================================================================
    # 1. CARD E80
    # =========================================================================
    add_heading_1(doc, "1. Card Điều Khiển Chuyển Động / Giao Tiếp E80 (Card E80)")
    add_paragraph(doc, "Card E80 là bo mạch điều khiển chuyển động và mở rộng I/O chuyên dụng trong công nghiệp, thường dùng làm cầu nối trung gian giữa máy tính điều khiển (qua USB/Ethernet/RS485) với các Driver động cơ bước/servo, cảm biến hành trình, công tắc an toàn và biến tần spindle.")
    
    add_heading_2(doc, "A. Chức năng & Công dụng Chính:")
    add_bullet(doc, "Phát xung điều khiển Step/Dir độc lập cho từ 3 đến 6 trục (X, Y, Z, A, B, C) với tần số xung cao, ổn định.", "• Điều khiển đa trục: ")
    add_bullet(doc, "Thu thập tín hiệu từ các cảm biến tiệm cận, công tắc hành trình Limit X/Y/Z, nút dừng khẩn cấp E-Stop, cảm biến Probe đo dao qua mạch cách ly quang Optocoupler.", "• Đọc tín hiệu Input cách ly: ")
    add_bullet(doc, "Đóng ngắt rơ-le bơm nước làm mát, hút bụi, xuất tín hiệu Analog 0-10V hoặc PWM để điều khiển tốc độ biến tần Spindle.", "• Điều khiển tải ngoại vi: ")
    add_bullet(doc, "Bảo vệ máy tính chủ khỏi xung sét cảm ứng, nhiễu động cơ và chập cháy điện áp cao trong môi trường nhà xưởng.", "• Cách ly công nghiệp: ")

    add_heading_2(doc, "B. Sơ đồ Chân & Cổng Kết Nối:")
    headers_e80 = ["Cụm Chân / Cổng", "Tên Ký Hiệu", "Chức Năng Chi Tiết", "Mức Điện Áp / Chuẩn"]
    data_e80 = [
        ["Cổng Nguồn Bo Mạch", "24V / GND (hoặc 5V)", "Cấp nguồn nuôi bo mạch điều khiển và mạch cách ly quang", "24VDC (Chuẩn CN) / 5VDC"],
        ["Cụm Xung Trục X", "XPUL+, XPUL-, XDIR+, XDIR-", "Ngõ ra xung bước và chiều quay nối sang Driver Trục X", "5V Differential / Open-Collector"],
        ["Cụm Xung Trục Y", "YPUL+, YPUL-, YDIR+, YDIR-", "Ngõ ra xung bước và chiều quay nối sang Driver Trục Y", "5V Differential / Open-Collector"],
        ["Cụm Xung Trục Z", "ZPUL+, ZPUL-, ZDIR+, ZDIR-", "Ngõ ra xung bước và chiều quay nối sang Driver Trục Z", "5V Differential / Open-Collector"],
        ["Cụm Xung Trục A", "APUL+, APUL-, ADIR+, ADIR-", "Ngõ ra trục xoay thứ 4 (Trục A / Rotary axis)", "5V Differential / Open-Collector"],
        ["Cụm Cảm Biến Vào (Input)", "IN1 -> IN8 (X-LIM, Y-LIM, Z-LIM, ESTOP, PROBE)", "Ngõ vào cách ly quang nối công tắc hành trình, nút dừng khẩn", "NPN / PNP (12V - 24VDC)"],
        ["Cụm Rơ-le / Ngoại Vi", "OUT1 -> OUT4, RELAY1 (COM/NO)", "Ngõ ra đóng van khí nén, bơm tưới nguội, đèn báo tháp", "Relay 250VAC 5A / 24VDC"],
        ["Điều Khiển Spindle", "AVI (0-10V), ACM (GND), FWD", "Xuất điện áp tuyến tính 0-10V điều khiển tốc độ biến tần VFD", "Analog 0-10V & Relay FWD"],
        ["Cổng Giao Tiếp Máy Chủ", "USB Type-B / RJ45 LAN / RS485", "Cắm cáp kết nối máy tính điều khiển (Mach3, LinuxCNC, App)", "Chuẩn giao tiếp số công nghiệp"]
    ]
    add_table_data(doc, headers_e80, data_e80, [1.4, 1.8, 2.4, 1.2])

    add_heading_2(doc, "C. Những Thứ Cần Lưu Ý Khi Sử Dụng:")
    add_bullet(doc, "Nguồn cấp cho mạch cách ly I/O (24VDC) phải lấy từ bộ nguồn độc lập hoặc qua lọc nguồn, không lấy chung nguồn với cuộn coil contactor lớn để tránh nhiễu.", "1. Phân tách nguồn cách ly: ")
    add_bullet(doc, "Tất cả vỏ kim loại máy CNC, vỏ spindle, vỏ card E80 và cọc PE của nguồn Mean Well phải được tiếp địa nối đất chung (Star Grounding) để triệt tiêu nhiễu cao tần.", "2. Tiếp địa chống nhiễu: ")
    add_bullet(doc, "Tuyệt đối không cấp nguồn ngược cực tính 24V vào cổng nguồn (cháy diode bảo vệ và tụ lọc nguồn).", "3. Cực tính nguồn: ")

    # =========================================================================
    # 2. DRIVER BEST BH57
    # =========================================================================
    add_heading_1(doc, "2. Driver Động Cơ Bước Vòng Kín Best BH57 (倍斯特智能)")
    add_paragraph(doc, "Driver Best BH57 là bộ điều khiển động cơ bước vòng kín 2 pha kỹ thuật số thế hệ mới (Closed-Loop Stepper Driver / Hybrid Servo), tích hợp vi xử lý DSP 32-bit tốc độ cao, chuyên dùng để kéo các dòng động cơ bước có gắn bộ mã hóa Encoder (size 57 hoặc 42).")

    add_heading_2(doc, "A. Chức năng & Công dụng Chính:")
    add_bullet(doc, "Nhận tín hiệu hồi tiếp vị trí từ Encoder ở đuôi động cơ với tốc độ micro-second, liên tục so sánh với lệnh phát xung từ ESP32/PLC.", "• Chống mất bước tuyệt đối: ")
    add_bullet(doc, "Tự động tăng dòng điện khi gặp tải nặng và hạ dòng điện khi chạy không tải hoặc dừng $\\rightarrow$ Động cơ chạy cực êm, mát và không bị gầm rú.", "• Điều khiển dòng biến thiên (Current Vector Control): ")
    add_bullet(doc, "Có ngõ ra báo động ALARM kích hoạt ngay khi kẹt trục hoặc đứt dây phản hồi Encoder, bảo vệ cơ khí không bị va đập gãy hỏng.", "• Tự ngắt bảo vệ an toàn: ")

    add_heading_2(doc, "B. Sơ đồ Chân & Cổng Kết Nối Chi Tiết:")
    headers_bh57 = ["Cụm Chân", "Ký Hiệu", "Chức Năng Chi Tiết", "Cách Đấu Nối Chuẩn"]
    data_bh57 = [
        ["Tín Hiệu Logic (5~24V)", "PUL+", "Cực dương tín hiệu xung bước", "Nối vào 3.3V của ESP32-S3 (Đấu Dương chung)"],
        ["", "PUL-", "Cực âm tín hiệu xung bước (Step)", "Nối vào GPIO 4 của ESP32-S3 (kéo LOW để phát xung)"],
        ["", "DIR+", "Cực dương tín hiệu chiều quay", "Nối vào 3.3V của ESP32-S3"],
        ["", "DIR-", "Cực âm tín hiệu chiều quay (Direction)", "Nối vào GPIO 5 của ESP32-S3 (LOW=Thuận, HIGH=Ngược)"],
        ["", "EN+", "Cực dương tín hiệu bật/tắt driver", "Nối vào 3.3V của ESP32-S3"],
        ["", "EN-", "Cực âm tín hiệu bật/tắt driver (Enable)", "Nối vào GPIO 6 hoặc BỎ TRỐNG (để trống = luôn BẬT)"],
        ["Báo Lỗi & Mở Rộng", "ALM+", "Cực dương ngõ ra cảnh báo lỗi", "Nối lên nguồn 3.3V"],
        ["", "ALM-", "Cực âm ngõ ra cảnh báo lỗi (Fault)", "Nối vào GPIO 7 (ESP32) để dừng khẩn khi kẹt tải"],
        ["", "EX+, EX-", "Cặp tín hiệu phụ trợ mở rộng", "Tùy biến qua phần mềm RS232 (bình thường để trống)"],
        ["Cáp Encoder Động Cơ", "VCC", "Nguồn dương +5V nuôi mắt đọc Encoder", "Nối dây ĐỎ của cáp Encoder"],
        ["", "EGND", "Nguồn âm 0V của Encoder", "Nối dây TRẮNG (hoặc Đen nhỏ) của cáp Encoder"],
        ["", "EB+, EB-", "Kênh xung B vi sai chống nhiễu", "Nối dây VÀNG (B+) và XANH LÁ (B-) của Encoder"],
        ["", "EA+, EA-", "Kênh xung A vi sai chống nhiễu", "Nối dây ĐEN TO (A+) và XANH DƯƠNG (A-) của Encoder"],
        ["Nguồn Động Lực", "V+", "Cực dương nguồn DC động lực", "Nối vào cọc V+ nguồn Mean Well 24VDC (hoặc 36V/48V)"],
        ["", "V-", "Cực âm nguồn DC động lực (GND)", "Nối vào cọc V- (COM) nguồn Mean Well 24VDC"],
        ["Cuộn Dây Động Cơ", "A+, A-", "Đầu cuộn dây Pha A của động cơ", "Nối dây ĐEN (A+) và XANH LÁ (A-) của động cơ"],
        ["", "B+, B-", "Đầu cuộn dây Pha B của động cơ", "Nối dây ĐỎ (B+) và XANH DƯƠNG (B-) của động cơ"],
        ["Cổng Phụ & Đèn", "RS232 / PWR / FLT", "Cổng tinh chỉnh tham số / LED Xanh nguồn / LED Đỏ lỗi", "PWR sáng xanh = tốt, FLT sáng đỏ = lỗi vị trí/quá tải"]
    ]
    add_table_data(doc, headers_bh57, data_bh57, [1.5, 1.2, 2.3, 1.8])

    add_heading_2(doc, "C. Những Thứ Cần Lưu Ý Khi Sử Dụng:")
    add_alert_box(
        doc,
        "⛔ CẢNH BÁO QUAN TRỌNG VỀ DRIVER BH57:",
        "1. KHÔNG CẮM ĐỘNG CƠ KHÔNG CÓ ENCODER: Driver BH57 bắt buộc phải nhận tín hiệu Encoder. Cắm động cơ thường sẽ bị nhảy đèn đỏ FLT ngắt nguồn ngay lập tức.\n"
        "2. NẾU ĐỘNG CƠ GIẬT MẠNH KHI VỪA BẬT: Do ngược kênh Encoder hoặc ngược pha cuộn dây. Cách xử lý: Đảo vị trí 2 dây A+ và A- cho nhau.\n"
        "3. CẤM RÚT CẮM DÂY MOTOR KHI ĐANG CÓ ĐIỆN 24V: Dòng xả cảm ứng ngược (Back-EMF) sẽ phá hủy tức thì cầu MOSFET công suất.",
        "danger"
    )

    # =========================================================================
    # 3. RASPBERRY PI 4 MODEL B (4GB)
    # =========================================================================
    add_heading_1(doc, "3. Máy Tính Nhúng Raspberry Pi 4 Model B (4GB RAM)")
    add_paragraph(doc, "Raspberry Pi 4 Model B là máy tính nhúng kích thước nhỏ mạnh mẽ (CPU Broadcom BCM2711 Quad-core Cortex-A72 1.5GHz, 4GB LPDDR4), đóng vai trò làm bộ xử lý trung tâm (Master Controller), chạy giao diện điều khiển HMI, Web Server, MQTT Broker hoặc các thuật toán xử lý ảnh thị giác máy tính.")

    add_heading_2(doc, "A. Chức năng & Công dụng Chính:")
    add_bullet(doc, "Chạy hệ điều hành Linux (Raspberry Pi OS, Ubuntu), xử lý đa nhiệm mượt mà, lưu trữ dữ liệu sản xuất vào thẻ nhớ MicroSD hoặc ổ cứng SSD.", "• Bộ xử lý trung tâm (Host Controller): ")
    add_bullet(doc, "Giao tiếp với ESP32-S3 qua cổng USB Serial hoặc UART, gửi lệnh vận hành và nhận trạng thái từ ESP32-S3.", "• Điều khiển & Giám sát: ")
    add_bullet(doc, "Mở rộng 2 cổng USB 3.0 (xanh) tốc độ cao cắm camera công nghiệp, 2 cổng USB 2.0 cắm chuột bàn phím, 2 cổng Micro HDMI xuất màn hình 4K, cổng Gigabit Ethernet.", "• Kết nối ngoại vi toàn diện: ")

    add_heading_2(doc, "B. Sơ đồ Chân Header 40 Pin GPIO:")
    headers_pi4 = ["Chân Vật Lý (Pin #)", "Tên Chân (Pin Name)", "Chức Năng & Ngoại Vi Hỗ Trợ", "Mức Điện Áp"]
    data_pi4 = [
        ["Pin 1, Pin 17", "3V3 Power", "Nguồn dương 3.3V cấp cho cảm biến công suất nhỏ (Max 50mA)", "3.3V DC Out"],
        ["Pin 2, Pin 4", "5V Power", "Nguồn dương 5V nối trực tiếp vào đường nguồn chính", "5V DC In/Out"],
        ["Pin 6, 9, 14, 20, 25, 30, 34, 39", "GND", "Mass đất chung hệ thống tín hiệu logic", "0V (Ground)"],
        ["Pin 3 (GPIO 2), Pin 5 (GPIO 3)", "SDA1 / SCL1", "Giao tiếp I2C bus chính (đã kéo trở nội lên 3.3V)", "Logic 3.3V"],
        ["Pin 8 (GPIO 14), Pin 10 (GPIO 15)", "TXD0 / RXD0", "Cổng UART truyền nhận nối tiếp tốc độ cao với ESP32", "Logic 3.3V"],
        ["Pin 12 (GPIO 18)", "PWM0", "Ngõ ra băm xung phần cứng PWM điều khiển tốc độ/servo", "Logic 3.3V"],
        ["Pin 19, 21, 23, 24, 26", "MOSI, MISO, SCLK, CE0, CE1", "Giao tiếp SPI bus tốc độ cao đọc cảm biến / màn hình phụ", "Logic 3.3V"],
        ["Pin 7, 11, 13, 15, 16, 18, 22...", "GPIO...", "Các chân I/O số lập trình tự do (Input, Output, Interrupt)", "Logic 3.3V"]
    ]
    add_table_data(doc, headers_pi4, data_pi4, [1.4, 1.5, 2.7, 1.2])

    add_heading_2(doc, "C. Những Thứ Cần Lưu Ý Khi Sử Dụng:")
    add_bullet(doc, "Tuyệt đối không cấp quá 3.3V vào bất kỳ chân GPIO nào (Pi 4 không có tính năng chịu áp 5V, cấp 5V vào GPIO sẽ chết chip BCM2711 ngay lập tức).", "1. Điện áp Logic GPIO 3.3V: ")
    add_bullet(doc, "Cần nguồn cấp 5V chuẩn dòng tối thiểu 3.0A (15W) qua cổng USB-C. Nếu nguồn yếu, Pi sẽ hiện biểu tượng tia sét vàng (Under-voltage) và giảm xung nhịp CPU.", "2. Nguồn cấp USB Type-C 5V/3A: ")
    add_bullet(doc, "Chip Pi 4 tỏa nhiệt lớn khi chạy tác vụ nặng. Bắt buộc lắp quạt tản nhiệt hoặc vỏ nhôm tản nhiệt để duy trì nhiệt độ dưới 65°C.", "3. Tản nhiệt bắt buộc: ")
    add_bullet(doc, "Luôn dùng lệnh 'sudo poweroff' hoặc 'sudo shutdown -h now' trước khi rút nguồn để tránh làm hỏng cấu trúc hệ điều hành trên thẻ nhớ SD.", "4. Tắt máy đúng quy trình: ")

    # =========================================================================
    # 4. MODULE HIỂN THỊ 8 LED & KEY TM1638
    # =========================================================================
    add_heading_1(doc, "4. Module Hiển Thị & Bàn Phím TM1638 (8 LED 7 Đoạn & 8 Phím Bấm)")
    add_paragraph(doc, "Module TM1638 (LED & KEY) là khối giao diện người dùng phần cứng tích hợp IC điều khiển quét phím và hiển thị TM1638, gồm 8 LED 7 đoạn, 8 đèn LED đơn đỏ và 8 nút nhấn, giao tiếp với ESP32-S3 chỉ bằng 3 đường tín hiệu.")

    add_heading_2(doc, "A. Chức năng & Công dụng Chính:")
    add_bullet(doc, "Hiển thị tốc độ (Hz), số bước đã quay, số vòng thực tế, mã lỗi hệ thống và menu cài đặt mà không cần màn hình cồng kềnh.", "• Hiển thị số liệu trực quan: ")
    add_bullet(doc, "Báo hiệu chiều quay (Thuận/Ngược), chế độ chạy liên tục, trạng thái Dừng/Khẩn cấp qua 8 đèn LED đơn đỏ D1 -> D8.", "• Đèn LED trạng thái: ")
    add_bullet(doc, "8 nút nhấn S1 -> S8 dùng để bấm quay 1 vòng, 2 vòng, quay 90 độ, tăng tốc, giảm tốc, đảo chiều, chạy liên tục (RUN) và dừng khẩn (STOP).", "• Bàn phím điều khiển thao tác: ")
    add_bullet(doc, "IC TM1638 tự động quét LED và chống rung phím (Debounce) bằng phần cứng, giúp ESP32-S3 không tốn tài nguyên CPU.", "• Tiết kiệm chân MCU: ")

    add_heading_2(doc, "B. Sơ đồ Chân Header J1 & Bảng Phím Bấm:")
    headers_tm = ["Chân Header J1", "Chức Năng", "Kết Nối Sang ESP32-S3", "Ghi Chú Kỹ Thuật"]
    data_tm = [
        ["VCC", "Nguồn dương nuôi bo mạch", "Nối chân 3.3V (hoặc 5V) của ESP32-S3", "Dòng tiêu thụ max ~130mA khi sáng full LED"],
        ["GND", "Mass âm nguồn (0V)", "Nối chân GND của ESP32-S3", "Nối chung mass với hệ thống"],
        ["STB (Strobe)", "Chân chọn chip (Chip Select)", "Nối vào GPIO 10 của ESP32-S3", "Mức LOW kích hoạt truyền nhận lệnh"],
        ["CLK (Clock)", "Chân xung nhịp đồng hồ", "Nối vào GPIO 11 của ESP32-S3", "Dữ liệu được đọc/ghi theo sườn xung CLK"],
        ["DIO (Data I/O)", "Chân dữ liệu 2 chiều", "Nối vào GPIO 12 của ESP32-S3", "Truyền mã LED xuất ra và đọc trạng thái 8 nút bấm về"]
    ]
    add_table_data(doc, headers_tm, data_tm, [1.3, 1.8, 2.0, 1.7])

    add_heading_2(doc, "C. Những Thứ Cần Lưu Ý Khi Sử Dụng:")
    add_bullet(doc, "Chân DIO là đường truyền 2 chiều (Half-duplex). Khi đọc phím, code trên ESP32 phải chuyển GPIO 12 sang chế độ INPUT, sau đó chuyển lại OUTPUT khi quét LED.", "1. Giao thức chân DIO: ")
    add_bullet(doc, "Module hoạt động hoàn hảo ở cả mức nguồn 3.3V và 5V. Khuyên dùng nguồn 3.3V trực tiếp từ ESP32-S3 để đồng bộ mức điện áp logic.", "2. Điện áp nuôi mạch: ")
    add_bullet(doc, "Khi nối dây dài > 30cm, nên dùng dây bọc giáp hoặc xoắn đôi các cặp dây CLK/GND, DIO/GND để tránh nhiễu xung làm hiển thị sai số.", "3. Chống nhiễu đường truyền: ")

    # =========================================================================
    # 5. HUB NGUỒN USB
    # =========================================================================
    add_heading_1(doc, "5. Hub Nguồn USB (Active / Industrial Powered USB Hub)")
    add_paragraph(doc, "Hub nguồn USB công nghiệp (Industrial Powered USB Hub) là bộ chia cổng và phân phối nguồn USB độc lập, có nguồn cấp DC ngoài riêng biệt, chuyên dùng để cấp nguồn và truyền dữ liệu cho nhiều thiết bị như Raspberry Pi 4, ESP32-S3, Camera công nghiệp, Màn hình cảm ứng mà không gây sụt áp cổng USB của máy tính chủ.")

    add_heading_2(doc, "A. Chức năng & Công dụng Chính:")
    add_bullet(doc, "Cung cấp dòng điện lớn (từ 2.4A đến 3.5A trên mỗi cổng) ổn định ở chuẩn 5.0V - 5.2V DC, không bị phụ thuộc vào dòng cấp yếu 500mA từ cổng máy tính.", "• Cấp nguồn công suất lớn: ")
    add_bullet(doc, "Truyền tín hiệu dữ liệu USB 2.0 / USB 3.0 tốc độ cao giữa Raspberry Pi / PC với nhiều mạch vi điều khiển ESP32, mạch nạp và cảm biến cùng lúc.", "• Mở rộng kết nối dữ liệu: ")
    add_bullet(doc, "Bảo vệ chống sốc tĩnh điện ESD (lên đến 15kV), chống quá dòng (OCP), chống ngắn mạch và ngăn dòng điện chạy ngược (Back-power/Back-feed) làm cháy cổng USB của Raspberry Pi/PC.", "• Cách ly bảo vệ an toàn: ")

    add_heading_2(doc, "B. Cổng Kết Nối & Sơ Đồ Chân Chuẩn USB:")
    headers_hub = ["Cổng / Chân", "Tên Tín Hiệu", "Mô Tả Chức Năng", "Thông Số Điện"]
    data_hub = [
        ["Cổng Nguồn Ngoài (DC In)", "DC Jack / Terminal Block", "Đầu vào cấp nguồn chính cho toàn bộ các cổng USB", "12V - 24V DC (hoặc 5V DC)"],
        ["Cổng Host (USB-B / Type-C)", "Upstream USB Port", "Cáp kết nối về máy tính chủ (Raspberry Pi 4 hoặc PC)", "Truyền dữ liệu Host"],
        ["Cổng Ra 1 -> 7 (USB-A)", "Downstream USB Ports", "Cấp nguồn và cắm thiết bị ngoại vi (ESP32-S3, Camera...)", "Chuẩn 5V DC (Max 2.4A/cổng)"],
        ["Chân 1 trên cổng USB-A", "VBUS (+5V)", "Dây cấp nguồn dương màu Đỏ", "5.0V DC (+/- 5%)"],
        ["Chân 2 trên cổng USB-A", "D- (Data Negative)", "Dây tín hiệu dữ liệu vi sai âm màu Trắng", "3.3V Logic USB"],
        ["Chân 3 trên cổng USB-A", "D+ (Data Positive)", "Dây tín hiệu dữ liệu vi sai dương màu Xanh lá", "3.3V Logic USB"],
        ["Chân 4 trên cổng USB-A", "GND (Ground)", "Dây nối mass nguồn màu Đen", "0V (Mass chung)"]
    ]
    add_table_data(doc, headers_hub, data_hub, [1.5, 1.4, 2.5, 1.4])

    add_heading_2(doc, "C. Những Thứ Cần Lưu Ý Khi Sử Dụng:")
    add_bullet(doc, "Tuyệt đối không dùng Hub USB không có nguồn phụ (Passive Hub) để nuôi Raspberry Pi hoặc nhiều vi điều khiển, vì sẽ gây sụt áp liên tục làm treo chip.", "1. Phải dùng Hub có nguồn phụ (Powered Hub): ")
    add_bullet(doc, "Chọn Hub có tính năng 'Anti-Backfeed' để dòng 5V từ Hub không xả ngược vào bo mạch chủ PC/Pi khi tắt nguồn.", "2. Chống dòng xả ngược: ")

    # =========================================================================
    # 6. VI ĐIỀU KHIỂN ESP32-S3
    # =========================================================================
    add_heading_1(doc, "6. Vi Điều Khiển ESP32-S3 (DevKitC-1 N16R8 / Dual Type-C)")
    add_paragraph(doc, "ESP32-S3 DevKitC-1 là bo mạch vi điều khiển 32-bit lõi kép Xtensa LX7 chạy xung nhịp lên tới 240MHz, tích hợp 16MB Flash, 8MB PSRAM, Wi-Fi 2.4GHz và Bluetooth 5.0 LE, chuyên đảm nhiệm việc phát xung điều khiển bước thời gian thực (FreeRTOS Multi-tasking) và xử lý giao diện người dùng.")

    add_heading_2(doc, "A. Chức năng & Công dụng Chính:")
    add_bullet(doc, "Tạo chuỗi xung bước siêu chính xác từ 10Hz đến 200kHz bằng phần cứng chuyên dụng (RMT / MCPWM / LEDC), không bị delay bởi tác vụ khác.", "• Phát xung điều khiển bước (Motion Engine): ")
    add_bullet(doc, "Bộ đếm xung phần cứng PCNT đọc tín hiệu Encoder A/B tốc độ hàng triệu xung/giây mà không tốn tài nguyên xử lý.", "• Đọc xung Encoder phần cứng (PCNT): ")
    add_bullet(doc, "Chạy hệ điều hành thời gian thực FreeRTOS, phân luồng riêng biệt: Luồng phát xung, Luồng quét phím TM1638, Luồng giao tiếp Serial.", "• Đa nhiệm đa luồng mượt mà: ")

    add_heading_2(doc, "B. Sơ đồ Phân Bổ Chân GPIO Đang Dùng Trong Dự Án:")
    headers_esp = ["Chân Trên ESP32-S3", "Thiết Bị Kết Nối", "Chức Năng Điều Khiển", "Mức Điện Áp Logic"]
    data_esp = [
        ["3.3V", "Driver DM542E / Best BH57", "Cấp nguồn Dương Chung (PUL+, DIR+, EN+)", "3.3V DC (Max 500mA)"],
        ["GND", "Toàn bộ hệ thống", "Mass âm đất chung của mạch logic", "0V"],
        ["GPIO 4", "PUL- trên Driver", "Ngõ ra xung bước (Phát xung kích mức LOW)", "3.3V Logic"],
        ["GPIO 5", "DIR- trên Driver", "Ngõ ra chiều quay (LOW=Thuận, HIGH=Nghịch)", "3.3V Logic"],
        ["GPIO 6", "ENA- trên Driver", "Ngõ ra bật/tắt động lực (Enable)", "3.3V Logic"],
        ["GPIO 7", "ALM- trên Driver", "Ngõ vào ngắt nhận tín hiệu báo lỗi / Kẹt tải", "3.3V Logic (Pull-up)"],
        ["GPIO 10", "STB trên TM1638", "Chân chọn chip Strobe Module LED & KEY", "3.3V Logic"],
        ["GPIO 11", "CLK trên TM1638", "Chân xung đồng hồ Clock quét LED", "3.3V Logic"],
        ["GPIO 12", "DIO trên TM1638", "Chân dữ liệu 2 chiều Data I/O đọc/ghi", "3.3V Logic"],
        ["GPIO 1, GPIO 2", "Màn hình LCD 20x4", "Chân I2C SDA (GPIO 1) và SCL (GPIO 2)", "3.3V Logic"],
        ["GPIO 48", "Đèn RGB Trên Bo Mạch", "Đèn LED đổi màu báo trạng thái chạy (Xanh/Đỏ/Vàng/Tím)", "Nội bộ trên bo mạch"]
    ]
    add_table_data(doc, headers_esp, data_esp, [1.4, 1.8, 2.4, 1.2])

    add_heading_2(doc, "C. Những Thứ Cần Lưu Ý & Cảnh Báo Sống Còn:")
    add_alert_box(
        doc,
        "⛔ CẢNH BÁO NGUY HIỂM VỀ ĐIỆN ÁP TRÊN ESP32-S3:",
        "1. ĐIỆN ÁP TỐI ĐA 3.3V: ESP32-S3 chỉ hoạt động ở 3.3V (KHÔNG chịu được 5V). Cấp 5V hoặc 24V vào bất kỳ chân GPIO nào sẽ làm nổ vi điều khiển ngay lập tức.\n"
        "2. CHÂN STRAPPING CẦN TRÁNH: GPIO 0, GPIO 45, GPIO 46, GPIO 3. Không được kéo cố định mức điện áp các chân này lúc khởi động vì sẽ làm ESP32 không boot được.\n"
        "3. CẤM DÙNG GPIO 26 ĐẾN GPIO 32: Đây là các chân kết nối chip nhớ Flash/PSRAM nội bộ Octal SPI.",
        "danger"
    )

    # =========================================================================
    # 7. MẠCH LỌC DẬP XUNG RC RD551150-270
    # =========================================================================
    add_heading_1(doc, "7. Mạch Lọc Dập Xung Điện Cảm Ứng RC RD551150-270 (RC Snubber Network)")
    add_paragraph(doc, "Mạch RC RD551150-270 là cụm linh kiện lọc dập xung điện áp cao (RC Snubber / Surge Absorber) gồm điện trở công suất R = 270Ω mắc nối tiếp với tụ điện cao áp C = 0.15μF (hoặc 0.1μF - 0.27μF / 250V-630V), chuyên dùng để dập tia lửa điện và triệt tiêu điện áp cảm ứng ngược (Back-EMF).")

    add_heading_2(doc, "A. Chức năng & Công dụng Chính:")
    add_bullet(doc, "Khi ngắt rơ-le, contactor, van điện từ khí nén (Solenoid) hoặc cuộn dây động cơ, hiện tượng cảm ứng điện từ sinh ra tia lửa hồ quang hàng nghìn Volts. Mạch RC sẽ hấp thụ toàn bộ năng lượng này.", "• Triệt tiêu tia lửa điện hồ quang: ")
    add_bullet(doc, "Xung nhiễu điện từ (EMI/RFI) phát ra từ cuộn dây khi đóng ngắt thường xuyên làm vi điều khiển ESP32 và Raspberry Pi bị đơ, treo hoặc reset ngẫu nhiên. Mạch RC dập tắt triệt để nguồn nhiễu này.", "• Chống treo / reset Vi điều khiển: ")
    add_bullet(doc, "Chống rỗ, dính tiếp điểm và cháy bề mặt kim loại của rơ-le / contactor, nâng cao tuổi thọ đóng ngắt lên gấp 10 lần.", "• Bảo vệ tiếp điểm Rơ-le: ")

    add_heading_2(doc, "B. Sơ đồ Đấu Nối & Thông Số:")
    headers_rc = ["Thông Số Kỹ Thuật", "Giá Trị Định Mức", "Vị Trí Đấu Nối Chuẩn", "Sơ Đồ Mắc"]
    data_rc = [
        ["Điện trở dập xung (R)", "270 Ohm (Công suất 2W - 5W)", "Mắc SONG SONG với 2 đầu cuộn dây Van điện từ / Contactor", "Mắc NỐI TIẾP R và C bên trong vỏ"],
        ["Tụ điện cao áp (C)", "0.15uF - 0.27uF / 275VAC (630VDC)", "Hoặc mắc SONG SONG với tiếp điểm COM-NO của Rơ-le", "Ra 2 đầu dây linh hoạt"],
        ["Điện áp làm việc", "AC 110V - 275V / DC 12V - 100V", "Càng đặt gần vị trí phát sinh tia lửa càng tốt", "Không phân biệt cực tính (+) (-)"]
    ]
    add_table_data(doc, headers_rc, data_rc, [1.6, 1.8, 2.0, 1.4])

    add_heading_2(doc, "C. Những Thứ Cần Lưu Ý Khi Sử Dụng:")
    add_bullet(doc, "Mạch RC có 2 đầu dây ra không phân cực tính (+/-), bạn có thể đấu đầu nào vào trước cũng được.", "1. Không phân biệt cực tính: ")
    add_bullet(doc, "Mạch phải được đấu SONG SONG trực tiếp ngay tại 2 cọc tiếp điểm hoặc 2 cọc cuộn dây phát sinh tia lửa (đấu càng gần nguồn sinh nhiễu hiệu quả càng cao).", "2. Đấu nối song song: ")
    add_bullet(doc, "Đối với tải thuần một chiều DC (như van 24VDC), có thể đấu kết hợp thêm 1 Diode dập xung ngược (như 1N4007 / SS34) phân cực ngược song song với cuộn dây để đạt hiệu quả dập xung 100%.", "3. Kết hợp Diode cho tải DC: ")

    # =========================================================================
    # 8. MÀN HÌNH LCD 20X4 (I2C)
    # =========================================================================
    add_heading_1(doc, "8. Màn Hình Ký Tự LCD 20x4 Kèm Module Giao Tiếp I2C (PCF8574)")
    add_paragraph(doc, "Màn hình LCD 20x4 (LCD 2004) là màn hình tinh thể lỏng hiển thị 4 dòng chữ, mỗi dòng tối đa 20 ký tự, được gắn sẵn bo mạch chuyển đổi I2C (dùng chip PCF8574) giúp giảm số chân kết nối với ESP32-S3 / Raspberry Pi từ 16 chân xuống chỉ còn 2 chân dữ liệu.")

    add_heading_2(doc, "A. Chức năng & Công dụng Chính:")
    add_bullet(doc, "Hiển thị đầy đủ cùng lúc 4 thông số: Tốc độ hiện tại (Hz), Tọa độ vị trí (Steps), Chiều quay và Trạng thái hệ thống (RUN / STOP / ALARM).", "• Hiển thị thông số vận hành: ")
    add_bullet(doc, "Hiển thị địa chỉ IP Wi-Fi, trạng thái kết nối MQTT/Web Server của ESP32-S3 hoặc Raspberry Pi.", "• Giám sát mạng & Kết nối: ")
    add_bullet(doc, "Có đèn nền xanh dương/xanh lá với biến trở tinh chỉnh độ đậm nhạt chữ ở mặt sau bo mạch.", "• Dễ dàng quan sát trong bóng tối: ")

    add_heading_2(doc, "B. Sơ đồ Chân Module I2C (4 Chân):")
    headers_lcd = ["Chân Module I2C", "Chức Năng", "Kết Nối ESP32-S3", "Kết Nối Raspberry Pi 4"]
    data_lcd = [
        ["GND", "Cực âm nguồn đất (0V)", "Nối chân GND", "Nối chân Pin 6 / Pin 9 (GND)"],
        ["VCC", "Nguồn dương cấp cho màn hình", "Nối chân 5V (hoặc VIN)", "Nối chân Pin 2 / Pin 4 (+5V)"],
        ["SDA (Serial Data)", "Đường truyền dữ liệu I2C", "Nối GPIO 1 của ESP32-S3", "Nối Pin 3 (GPIO 2 - SDA1)"],
        ["SCL (Serial Clock)", "Đường xung nhịp I2C", "Nối GPIO 2 của ESP32-S3", "Nối Pin 5 (GPIO 3 - SCL1)"]
    ]
    add_table_data(doc, headers_lcd, data_lcd, [1.3, 1.8, 1.9, 1.8])

    add_heading_2(doc, "C. Những Thứ Cần Lưu Ý Khi Sử Dụng:")
    add_bullet(doc, "Màn hình LCD 20x4 yêu cầu nguồn VCC = 5V thì chữ mới hiển thị đậm nét và đèn nền đủ sáng (nếu cấp 3.3V chữ sẽ rất mờ).", "1. Bắt buộc cấp nguồn VCC 5V: ")
    add_bullet(doc, "Nếu cấp nguồn mà màn hình chỉ hiện một hàng ô vuông đen hoặc không thấy chữ, hãy dùng tua-vít vặn biến trở chiết áp màu xanh ở mặt sau module I2C để chỉnh độ tương phản (Contrast).", "2. Chỉnh biến trở tương phản: ")
    add_bullet(doc, "Địa chỉ I2C mặc định trong code thường là '0x27' (với chip PCF8574T) hoặc '0x3F' (với chip PCF8574AT).", "3. Địa chỉ I2C Scanner: ")

    # =========================================================================
    # 9. ĐỘNG CƠ BƯỚC 42CM06-RD10
    # =========================================================================
    add_heading_1(doc, "9. Động Cơ Bước 42CM06-RD10 (Kèm Hộp Số Giảm Tốc 1:10)")
    add_paragraph(doc, "42CM06-RD10 là tổ hợp động cơ bước 2 pha chuẩn NEMA 17 (kích thước mặt bích 42x42mm, góc bước cơ bản 1.8°) được tích hợp sẵn cụm hộp số giảm tốc bánh răng hành tinh (Planetary Gearbox) tỷ số truyền 1:10 ở đầu trục.")

    add_heading_2(doc, "A. Chức năng & Công dụng Chính:")
    add_bullet(doc, "Hộp số 1:10 khuếch đại lực kéo mô-men xoắn ở đầu trục ra lên gấp 10 lần so với động cơ nguyên bản, giúp kéo tải nặng mà không bị đuối lực.", "• Tăng lực xoắn gấp 10 lần (Torque Multiplier): ")
    add_bullet(doc, "Để quay được 1 vòng trục ra của hộp số, động cơ phải quay 10 vòng $\\rightarrow$ Với vi bước 1600 xung/vòng, 1 vòng trục ra cần 16,000 xung $\\rightarrow$ Góc quay siêu mịn và siêu chính xác.", "• Tăng độ phân giải góc bước: ")
    add_bullet(doc, "Hộp số có độ rơ khe hở bánh răng (Backlash) cực nhỏ, cấu trúc kim loại chịu tải rung lắc cơ khí lớn.", "• Khóa trục chống trôi quán tính: ")

    add_heading_2(doc, "B. Bảng Quy Chuẩn 4 Dây Động Lực Động Cơ:")
    headers_m42 = ["Ký Hiệu Chân Driver", "Chức Năng Cuộn Dây", "Màu Dây Trên Động Cơ", "Đo Điện Trở (VOM)"]
    data_m42 = [
        ["A+", "Đầu dương cuộn dây Pha A", "Dây ĐEN (Black)", "Thông mạch với dây Xanh lá (~1.2 - 2.5 Ohm)"],
        ["A-", "Đầu âm cuộn dây Pha A", "Dây XANH LÁ (Green)", "Thông mạch với dây Đen"],
        ["B+", "Đầu dương cuộn dây Pha B", "Dây ĐỎ (Red)", "Thông mạch với dây Xanh dương (~1.2 - 2.5 Ohm)"],
        ["B-", "Đầu âm cuộn dây Pha B", "Dây XANH DƯƠNG (Blue)", "Thông mạch với dây Đỏ"]
    ]
    add_table_data(doc, headers_m42, data_m42, [1.4, 1.8, 1.8, 1.8])

    add_heading_2(doc, "C. Những Thứ Cần Lưu Ý Khi Sử Dụng:")
    add_alert_box(
        doc,
        "⚠️ CẢNH BÁO TƯƠNG THÍCH ĐỘNG CƠ 42CM06-RD10:",
        "1. KHÔNG CÓ ENCODER: Con động cơ này là động cơ bước vòng hở. Tương thích hoàn hảo với Driver Leadshine DM542E hoặc TB6600. KHÔNG dùng với Driver Best BH57 (sẽ báo lỗi FLT do thiếu Encoder).\n"
        "2. TÍNH TOÁN TỐC ĐỘ: Do có hộp số 1:10, khi muốn trục ra quay 1 vòng/giây (60 RPM) ở vi bước 1600, bạn phải phát tần số xung là 1600 x 10 = 16,000 Hz.\n"
        "3. CẤM RÚT DÂY MOTOR KHI CẤP ĐIỆN: Rút dây khi driver đang giữ dòng sẽ đánh thủng mạch công suất.",
        "danger"
    )

    # =========================================================================
    # 10. BỘ 3 NGUỒN MEAN WELL CÔNG NGHIỆP
    # =========================================================================
    add_heading_1(doc, "10. Bộ 3 Nguồn Tổ Ong Công Nghiệp Mean Well (LRS-50-5 / LRS-150-12 / LRS-100N2-24)")
    add_paragraph(doc, "Hệ thống sử dụng bộ 3 nguồn xung Mean Well chính hãng chuẩn công nghiệp (LRS Series) cung cấp các mức điện áp chuyên biệt 5V, 12V và 24V, đảm bảo hệ thống vận hành bền bỉ 24/7.")

    add_heading_2(doc, "A. Bảng Phân Công Nhiệm Vụ 3 Bộ Nguồn:")
    headers_psu = ["Mã Nguồn Mean Well", "Điện Áp & Dòng Max", "Công Suất", "Nhiệm Vụ Cấp Nguồn Trong Hệ Thống"]
    data_psu = [
        ["Mean Well LRS-50-5", "5.0V DC - 10.0A", "50 Watt", "Nuôi Raspberry Pi 4 (5V/3A), ESP32-S3, Module TM1638, LCD 20x4, USB Hub"],
        ["Mean Well LRS-150-12", "12.0V DC - 12.5A", "150 Watt", "Nuôi van điện từ khí nén 12V, quạt tản nhiệt tủ điện, cảm biến tiệm cận 12V, rơ-le"],
        ["Mean Well LRS-100N2-24", "24.0V DC - 4.5A (Peak 200% = 9A)", "100 Watt (Peak 200W)", "Cấp nguồn động lực công suất cho Driver DM542E / BH57 kéo động cơ bước"]
    ]
    add_table_data(doc, headers_psu, data_psu, [1.6, 1.4, 1.1, 2.7])

    add_heading_2(doc, "B. Sơ đồ Chân Đấu Nối Chuẩn Trên Nguồn Mean Well:")
    headers_term_psu = ["Tên Cọc Nối", "Ký Hiệu", "Mô Tả Chức Năng", "Quy Chuẩn Đấu Dây"]
    data_term_psu = [
        ["Pha Nóng AC 220V", "L (Line)", "Đầu vào điện lưới xoay chiều 220VAC", "Dây Đỏ hoặc Nâu (Qua cầu chì/CB)"],
        ["Pha Nguội AC 220V", "N (Neutral)", "Đầu vào điện lưới trung tính 220VAC", "Dây Xanh dương hoặc Đen"],
        ["Tiếp Địa Bảo Vệ", "PE / ⏚ (Ground)", "Chân nối đất vỏ kim loại chống rò giật", "Dây Vàng-Sọc Xanh nối vào vỏ tủ"],
        ["Cực Âm Nguồn DC", "-V / COM", "Cực âm nguồn một chiều (0VDC)", "Dây Đen (Tiết diện >= 0.75mm²)"],
        ["Cực Dương Nguồn DC", "+V", "Cực dương nguồn một chiều", "Dây Đỏ (Tiết diện >= 0.75mm²)"],
        ["Chiết Áp Tinh Chỉnh", "+V ADJ", "Biến trở xoay chỉnh điện áp ra (+/- 10%)", "Dùng tua-vít chỉnh chuẩn 5.1V, 12.0V, 24.0V"]
    ]
    add_table_data(doc, headers_term_psu, data_term_psu, [1.4, 1.2, 2.3, 1.9])

    add_heading_2(doc, "C. Những Thứ Cần Lưu Ý & Cảnh Báo An Toàn:")
    add_bullet(doc, "Bắt buộc nối cọc PE (tiếp địa) của cả 3 bộ nguồn vào vỏ kim loại của tủ điện để triệt tiêu điện rò cảm ứng và chống giật khi chạm vào máy.", "1. Tiếp địa vỏ nguồn (PE): ")
    add_bullet(doc, "Phải nối chung toàn bộ cọc COM (-V) của nguồn 5V, nguồn 12V và nguồn 24V về một thanh đồng Mass chung (Common Ground) để đồng pha mức 0V cho toàn bộ tín hiệu điều khiển.", "2. Nối chung Mass (Common Ground): ")
    add_bullet(doc, "Tuyệt đối không được cắm nhầm nguồn 24V vào đường nguồn 5V của Raspberry Pi hay ESP32 (cháy toàn bộ linh kiện điện tử ngay lập tức).", "3. Phân biệt rõ đường dây 5V và 24V: ")

    # =========================================================================
    # 11. ĐỘNG CƠ BƯỚC CÓ ENCODER
    # =========================================================================
    add_heading_1(doc, "11. Động Cơ Bước Vòng Kín Kèm Bộ Mã Hóa Encoder (Closed-Loop Stepper)")
    add_paragraph(doc, "Động cơ bước vòng kín (ví dụ: Leadshine 57CME13 / 57CME23 / JMC 42HSE) gồm thân động cơ bước 2 pha tích hợp sẵn cụm mắt đọc Encoder quang/từ tính độ phân giải cao (1000 xung/vòng = 4000 trạng thái đếm) được bảo vệ bằng nắp kim loại ở đuôi động cơ.")

    add_heading_2(doc, "A. Chức năng & Công dụng Chính:")
    add_bullet(doc, "Phát tín hiệu phản hồi góc quay thực tế về Driver Best BH57 hoặc Leadshine CL57T liên tục trong thời gian thực.", "• Hồi tiếp vị trí chính xác: ")
    add_bullet(doc, "Khi bị kẹt cơ khí hoặc tải đột ngột tăng cao, Encoder báo về để Driver tự động dồn 100% dòng điện kéo vượt qua vị trí kẹt $\\rightarrow$ Triệt tiêu hiện tượng trượt bước.", "• Tự động chống mất bước: ")
    add_bullet(doc, "Không bị giật cục và rung lắc cộng hưởng ở tốc độ thấp, vận hành êm ái tiệm cận động cơ AC Servo đắt tiền.", "• Vận hành êm ái: ")

    add_heading_2(doc, "B. Sơ đồ 2 Chùm Dây Kết Nối:")
    headers_enc_m = ["Chùm Dây", "Chân Driver BH57 / CL57T", "Màu Dây Quy Chuẩn", "Ý Nghĩa Kỹ Thuật"]
    data_enc_m = [
        ["Chùm 1: Động Lực (4 dây to)", "A+", "ĐEN (Black)", "Đầu dương cuộn dây Pha A động cơ"],
        ["", "A-", "XANH LÁ (Green)", "Đầu âm cuộn dây Pha A động cơ"],
        ["", "B+", "ĐỎ (Red)", "Đầu dương cuộn dây Pha B động cơ"],
        ["", "B-", "XANH DƯƠNG (Blue)", "Đầu âm cuộn dây Pha B động cơ"],
        ["Chùm 2: Encoder (6 dây nhỏ)", "VCC", "ĐỎ (Red)", "Nguồn dương +5V do Driver cấp nuôi Encoder"],
        ["", "EGND", "TRẮNG (White) / Đen nhỏ", "Mass âm nguồn 0V của Encoder"],
        ["", "EA+", "ĐEN (Black) / Nâu", "Kênh xung A pha dương"],
        ["", "EA-", "XANH DƯƠNG (Blue)", "Kênh xung A pha âm (Tín hiệu vi sai)"],
        ["", "EB+", "VÀNG (Yellow)", "Kênh xung B pha dương"],
        ["", "EB-", "XANH LÁ (Green)", "Kênh xung B pha âm (Tín hiệu vi sai)"]
    ]
    add_table_data(doc, headers_enc_m, data_enc_m, [1.6, 1.4, 1.8, 2.0])

    add_heading_2(doc, "C. Những Thứ Cần Lưu Ý & Khắc Phục Lỗi Ngược Pha:")
    add_alert_box(
        doc,
        "⚠️ LỖI KINH ĐIỂN KHI ĐẤU ĐỘNG CƠ CÓ ENCODER:",
        "• HIỆN TƯỢNG: Vừa cấp lệnh chạy, động cơ giật giật, rung bần bật 1-2 giây rồi Driver ngắt động lực, đèn đỏ FLT sáng lên báo lỗi 'Position Error'.\n"
        "• NGUYÊN NHÂN: Đấu ngược chiều pha cuộn dây hoặc ngược kênh A/B của Encoder (hồi tiếp dương làm driver bù sai hướng).\n"
        "• CÁCH XỬ LÝ CHUẨN 100%: Tắt nguồn điện, đổi chỗ 2 dây A+ và A- cho nhau (giữ nguyên pha B và cáp Encoder), bật lại là động cơ sẽ chạy êm ru!",
        "warning"
    )

    # Save to file
    try:
        doc.save(docx_path)
        print(f"Successfully created {docx_path}")
    except PermissionError:
        alt_path = docx_path.replace(".docx", "_NEW.docx")
        doc.save(alt_path)
        print(f"File was locked, saved to {alt_path}")

if __name__ == "__main__":
    generate_full_hardware_docx(r"d:\Luu\THIET_BI_PHAN_CUNG_CHI_TIET.docx")
