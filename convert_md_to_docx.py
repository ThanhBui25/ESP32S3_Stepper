import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def format_text_runs(paragraph, text, default_bold=False, default_italic=False, default_color=None, font_name="Segoe UI", font_size=10.5):
    tokens = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for token in tokens:
        if not token:
            continue
        run = paragraph.add_run()
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if default_color:
            run.font.color.rgb = default_color
            
        if token.startswith('***') and token.endswith('***') and len(token) >= 6:
            run.text = token[3:-3]
            run.bold = True
            run.italic = True
        elif token.startswith('**') and token.endswith('**') and len(token) >= 4:
            run.text = token[2:-2]
            run.bold = True
            run.italic = default_italic
        elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
            run.text = token[1:-1]
            run.bold = default_bold
            run.italic = True
        elif token.startswith('`') and token.endswith('`') and len(token) >= 2:
            run.text = token[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(font_size - 0.5)
            run.font.color.rgb = RGBColor(199, 37, 78) # pink/red code
            run.bold = True
        else:
            run.text = token
            run.bold = default_bold
            run.italic = default_italic

def insert_image_card(doc, img_path, caption):
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(2)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(6.2))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(1)
        p_cap.paragraph_format.space_after = Pt(8)
        r_cap = p_cap.add_run(f"📷 {caption}")
        r_cap.font.name = "Segoe UI"
        r_cap.font.size = Pt(9)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(100, 116, 139)

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = docx.Document()
    
    # Configure page margins
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
        # Header / Footer
        footer = s.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Tài liệu Phần cứng & Tương thích - ESP32S3 & Stepper Drivers")
        f_run.font.name = "Segoe UI"
        f_run.font.size = Pt(8.5)
        f_run.font.color.rgb = RGBColor(156, 163, 175)

    i = 0
    in_code_block = False
    code_block_lines = []
    
    # Image mapping
    image_rules = [
        ("CẤU HÌNH PHẦN CỨNG ĐANG HOẠT ĐỘNG", r"d:\Luu\images\tm1638_card.png", "Hình 0: Module Bàn phím & Màn hình hiển thị TM1638 kết nối ESP32-S3"),
        ("2.1. ESP32-S3", r"d:\Luu\images\esp32s3_card.png", "Hình 1: Thẻ nhận diện & Sơ đồ chân Vi điều khiển ESP32-S3 DevKitC-1"),
        ("3.1. Driver Leadshine DM542", r"d:\Luu\images\dm542e_card.png", "Hình 2: Thẻ nhận diện & Cổng kết nối Driver Leadshine DM542E"),
        ("3.4.2. Driver Vòng Kín Best BH57", r"d:\Luu\images\bh57_card.png", "Hình 3: Thẻ nhận diện & Cổng kết nối Driver Vòng Kín Best BH57 (倍斯特智能)"),
        ("3.5. Driver Trinamic TMC2209", r"d:\Luu\images\tmc2209_card.png", "Hình 4: Thẻ nhận diện & Sơ đồ chân Module Driver Trinamic TMC2209"),
        ("4.1. Động cơ bước 2 pha vòng hở", r"d:\Luu\images\motor_gearbox_card.png", "Hình 5: Động cơ bước 42CM06-RD kèm Hộp số giảm tốc & Bảng mã màu 4 dây"),
        ("4.2. Động cơ bước vòng kín kèm Encoder", r"d:\Luu\images\closed_loop_motor_card.png", "Hình 6: Động cơ bước Vòng Kín 57CME kèm Bộ mã hóa Optical Encoder")
    ]
    
    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        
        # Code block handling
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                # Write code block table
                code_text = '\n'.join(code_block_lines)
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                
                cell = table.cell(0, 0)
                cell.width = Inches(6.8)
                set_cell_background(cell, "F3F4F6")
                set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(31, 41, 55)
                
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
                code_block_lines = []
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue
            
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
            
        # Ignore horizontal rules
        if re.match(r'^\s*---\s*$', line):
            i += 1
            continue
            
        # Empty line
        if not line.strip():
            i += 1
            continue
            
        # Title (# )
        if line.startswith('# '):
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(title_text)
            run.font.name = "Segoe UI"
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138) # Navy
            i += 1
            continue
            
        # H2 (## )
        if line.startswith('## '):
            h2_text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(h2_text)
            run.font.name = "Segoe UI"
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(29, 78, 216) # Blue
            
            # Check image trigger for H2
            for key, img_p, cap in image_rules:
                if key in h2_text:
                    insert_image_card(doc, img_p, cap)
            i += 1
            continue
            
        # H3 (### )
        if line.startswith('### '):
            h3_text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(h3_text)
            run.font.name = "Segoe UI"
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(31, 41, 55) # Dark Grey
            
            # Check image trigger for H3
            for key, img_p, cap in image_rules:
                if key in h3_text:
                    insert_image_card(doc, img_p, cap)
            i += 1
            continue
            
        # H4 (#### )
        if line.startswith('#### '):
            h4_text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(h4_text)
            run.font.name = "Segoe UI"
            run.font.size = Pt(10.5)
            run.bold = True
            run.font.color.rgb = RGBColor(55, 65, 81)
            i += 1
            continue
            
        # Blockquote / Alert (> )
        if line.startswith('> '):
            quote_lines = []
            while i < len(lines) and (lines[i].startswith('> ') or lines[i].strip() == '>'):
                quote_lines.append(lines[i].lstrip('> ').strip())
                i += 1
            
            quote_text = ' '.join(quote_lines)
            
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            cell = table.cell(0, 0)
            cell.width = Inches(6.8)
            
            if 'CẢNH BÁO' in quote_text or '⛔' in quote_text or '⚠️' in quote_text:
                set_cell_background(cell, "FEF2F2")
                border_color = "DC2626"
            elif '📌' in quote_text or '💡' in quote_text:
                set_cell_background(cell, "EFF6FF")
                border_color = "3B82F6"
            else:
                set_cell_background(cell, "F9FAFB")
                border_color = "9CA3AF"
                
            set_cell_margins(cell, top=100, bottom=100, left=180, right=150)
            
            tcPr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
                f'  <w:top w:val="none"/>'
                f'  <w:bottom w:val="none"/>'
                f'  <w:right w:val="none"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(borders)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            format_text_runs(p, quote_text, font_size=10)
            
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue
            
        # Markdown Table Detection (| ... |)
        if line.startswith('|') and '|' in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
                
            if len(table_lines) >= 2:
                rows_data = []
                for tl in table_lines:
                    if re.match(r'^\|[\s\:\-\|]+$', tl):
                        continue
                    cells = [c.strip() for c in tl.strip('|').split('|')]
                    rows_data.append(cells)
                    
                if rows_data:
                    num_cols = max(len(r) for r in rows_data)
                    for r in rows_data:
                        while len(r) < num_cols:
                            r.append('')
                            
                    table = doc.add_table(rows=len(rows_data), cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.autofit = False
                    set_table_borders(table)
                    
                    col_width = Inches(6.8 / num_cols)
                    
                    for r_idx, r_data in enumerate(rows_data):
                        row = table.rows[r_idx]
                        is_header = (r_idx == 0)
                        
                        for c_idx, cell_value in enumerate(r_data):
                            cell = row.cells[c_idx]
                            cell.width = col_width
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                            
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(2)
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.line_spacing = 1.05
                            
                            if is_header:
                                set_cell_background(cell, "1E3A8A")
                                format_text_runs(p, cell_value, default_bold=True, default_color=RGBColor(255, 255, 255), font_size=9.5)
                            else:
                                if r_idx % 2 == 0:
                                    set_cell_background(cell, "F9FAFB")
                                else:
                                    set_cell_background(cell, "FFFFFF")
                                format_text_runs(p, cell_value, font_size=9)
                    
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue
            
        # Bullet list item (- or *)
        if re.match(r'^\s*[\-\*]\s+', line):
            indent_level = len(re.match(r'^\s*', line).group(0)) // 2
            item_text = re.sub(r'^\s*[\-\*]\s+', '', line)
            
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            format_text_runs(p, item_text, font_size=10.5)
            i += 1
            continue
            
        # Numbered list item (1. 2. etc)
        if re.match(r'^\s*\d+\.\s+', line):
            item_text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            format_text_runs(p, item_text, font_size=10.5)
            i += 1
            continue
            
        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        format_text_runs(p, line, font_size=10.5)
        i += 1
        
    try:
        doc.save(docx_path)
        print(f"Successfully converted {md_path} -> {docx_path} with images embedded.")
    except PermissionError:
        fallback_path = r"d:\Luu\HARDWARE_DATABASE_IMAGES.docx"
        doc.save(fallback_path)
        print(f"File {docx_path} is currently open in Word. Saved to {fallback_path} instead.")

if __name__ == "__main__":
    convert_md_to_docx(r"d:\Luu\HARDWARE_DATABASE.md", r"d:\Luu\HARDWARE_DATABASE.docx")

